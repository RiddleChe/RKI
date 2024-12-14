from fastapi import FastAPI, UploadFile, Form, HTTPException
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from pathlib import Path
import shutil
import uuid
from docx import Document
from openai import OpenAI

# Инициализация FastAPI
app = FastAPI()

# Папки для загруженных и обработанных файлов
UPLOAD_FOLDER = Path("uploads")
PROCESSED_FOLDER = Path("processed")
UPLOAD_FOLDER.mkdir(exist_ok=True)
PROCESSED_FOLDER.mkdir(exist_ok=True)

# Инициализация OpenAI API клиента
client = OpenAI(
    api_key="api_tocken",
    base_url="https://api.proxyapi.ru/openai/v1",
)

# Функция для извлечения текста из DOCX файла
def extract_text_from_docx(file_path: Path) -> str:
    try:
        doc = Document(file_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ошибка при чтении файла: {str(e)}")

# Функция для создания нового DOCX файла с текстом
def create_docx_with_text(text: str, output_path: Path):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line.strip())
    doc.save(output_path)

# Функция для обработки текста с помощью GPT
def adapt_text_with_gpt(text: str, level: str) -> str:
    prompt = [
        {"role": "system", "content": "Ты эксперт по адаптации текстов на русском языке под уровни A1, A2, B1, B2, C1, C2."},
        {"role": "user", "content": f"Адаптируй следующий текст под уровень {level}: {text}"}
    ]
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=prompt,
            temperature=0.7,
            max_tokens=1500,
            top_p=1.0
        )
        return response.choices[0].message.content
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при обработке текста через GPT: {str(e)}")

# Обработчик для загрузки и обработки файлов
@app.post("/upload/")
async def upload_and_process_file(file: UploadFile = Form(...), level: str = Form(...)):
    # Проверяем формат файла
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Только файлы формата DOCX поддерживаются")

    # Генерируем уникальное имя для файла
    unique_filename = f"{uuid.uuid4()}_{file.filename}"
    saved_file_path = UPLOAD_FOLDER / unique_filename

    # Сохраняем загруженный файл
    with saved_file_path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Извлечение текста из файла
    extracted_text = extract_text_from_docx(saved_file_path)
    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="Файл не содержит текста для обработки")

    # Обработка текста через GPT
    processed_text = adapt_text_with_gpt(extracted_text, level)

    # Создание нового обработанного DOCX файла
    processed_file_path = PROCESSED_FOLDER / unique_filename
    create_docx_with_text(processed_text, processed_file_path)

    # Возвращаем обработанный файл
    return {
        "message": "Файл успешно обработан",
        "file_url": f"/processed/{unique_filename}"
    }

# Подключаем папку processed как статические файлы
app.mount("/processed", StaticFiles(directory="processed"), name="processed")

# Подключаем статический фронтенд
app.mount("/", StaticFiles(directory="frontend", html=True), name="frontend")
